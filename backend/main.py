from fastapi import FastAPI, File, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from backend.core.orchestrator import Orchestrator
from backend.utils.parser import DocumentParser
from pydantic import BaseModel
from typing import List, Dict, Optional
import io

app = FastAPI(title="IDAS Project Backend")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Initialize Orchestrator
orchestrator = Orchestrator()

class AnalysisRequest(BaseModel):
    text: str

class RewriteRequest(BaseModel):
    text: str
    issue: str
    suggestion: str

@app.get("/")
def read_root():
    return {"message": "IDAS Backend is running"}

@app.post("/analyze")
def analyze_requirement(request: AnalysisRequest):
    result = orchestrator.process(request.text, task_type="analyze")
    return {"result": result}

@app.post("/check-template")
def check_template(request: AnalysisRequest):
    result = orchestrator.process(request.text, task_type="check_template")
    return {"result": result}

@app.post("/check-glossary")
def check_glossary(request: AnalysisRequest):
    result = orchestrator.process(request.text, task_type="check_glossary")
    return {"result": result}

@app.post("/draft")
def draft_requirements(request: AnalysisRequest):
    result = orchestrator.process(request.text, task_type="draft")
    return {"result": result}

@app.post("/classify")
def classify_requirement(request: AnalysisRequest):
    result = orchestrator.process(request.text, task_type="classify")
    return {"result": result}

@app.post("/audit")
def audit_requirement(request: AnalysisRequest):
    result = orchestrator.process(request.text, task_type="audit")
    return {"result": result}

@app.post("/resolve")
def resolve_issue(request: AnalysisRequest):
    result = orchestrator.process(request.text, task_type="resolve")
    return {"result": result}

@app.get("/templates")
def get_templates():
    import json
    import os
    current_dir = os.path.dirname(os.path.abspath(__file__))
    templates_path = os.path.join(current_dir, "data", "templates.json")
    with open(templates_path, "r", encoding="utf-8") as f:
        return json.load(f)

@app.get("/session")
def get_session():
    import json
    import os
    current_dir = os.path.dirname(os.path.abspath(__file__))
    session_path = os.path.join(current_dir, "data", "session.json")
    with open(session_path, "r", encoding="utf-8") as f:
        return json.load(f)

@app.post("/session")
def save_session(data: dict):
    import json
    import os
    current_dir = os.path.dirname(os.path.abspath(__file__))
    session_path = os.path.join(current_dir, "data", "session.json")
    with open(session_path, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
    return {"status": "success"}

@app.get("/ui-config")
def get_ui_config():
    import json
    import os
    current_dir = os.path.dirname(os.path.abspath(__file__))
    config_path = os.path.join(current_dir, "data", "ui_config.json")
    with open(config_path, "r", encoding="utf-8") as f:
        return json.load(f)

@app.post("/chat")
def chat_with_ai(request: AnalysisRequest):
    result = orchestrator.process(request.text, task_type="chat")
    return {"result": result}

@app.post("/upload-review")
async def upload_review(file: UploadFile = File(...)):
    try:
        content = await file.read()
        parsed_text = DocumentParser.parse(file.filename, content)
        result = orchestrator.process(parsed_text, task_type="review")
        return result
    except Exception as e:
        return {"error": str(e)}

@app.post("/rewrite")
def rewrite_requirement(request: RewriteRequest):
    result = orchestrator.process(
        request.text, 
        task_type="rewrite", 
        issue=request.issue, 
        suggestion=request.suggestion
    )
    return {"result": result}



# ── Export Models ──────────────────────────────────────────────────────────────

class RevisionRow(BaseModel):
    version: Optional[str] = "-"
    date: Optional[str] = "-"
    name: Optional[str] = "-"
    reason: Optional[str] = "-"

class SectionMeta(BaseModel):
    id: str
    title: str

class ExportDocxRequest(BaseModel):
    doc_title: str
    sections: List[SectionMeta]
    content: Dict[str, str]
    revision_history: List[RevisionRow]


@app.post("/export-docx")
def export_docx(request: ExportDocxRequest):
    """Generate a real .docx file using python-docx and stream it as download."""
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import datetime

    doc = Document()

    # Page margins
    for sec in doc.sections:
        sec.top_margin    = Inches(1)
        sec.bottom_margin = Inches(1)
        sec.left_margin   = Inches(1.25)
        sec.right_margin  = Inches(1.25)

    # Title
    title_para = doc.add_heading(request.doc_title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title_para.runs:
        run.font.size = Pt(20)
        run.font.bold = True

    doc.add_paragraph()

    # Revision History Table
    doc.add_heading("Revision History", level=2)
    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    for i, col_name in enumerate(["Version", "Date", "Author", "Reason"]):
        hdr[i].text = col_name
        hdr[i].paragraphs[0].runs[0].font.bold = True

    for row in request.revision_history:
        cells = tbl.add_row().cells
        cells[0].text = row.version or "-"
        cells[1].text = row.date    or "-"
        cells[2].text = row.name   or "-"
        cells[3].text = row.reason or "-"

    doc.add_paragraph()

    # Sections
    for idx, sec in enumerate(request.sections, start=1):
        doc.add_heading(f"{idx}. {sec.title}", level=1)
        text = request.content.get(sec.id, "").strip() or "This section has not been filled out yet."
        for line in text.split("\n"):
            doc.add_paragraph(line)

    # Footer
    footer_para = doc.sections[0].footer.paragraphs[0]
    footer_para.text = f"Generated by IDAS  •  {datetime.date.today().strftime('%B %d, %Y')}"
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Stream as .docx
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    safe_name = request.doc_title.replace(" ", "_").replace("/", "-")
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{safe_name}.docx"'},
    )
